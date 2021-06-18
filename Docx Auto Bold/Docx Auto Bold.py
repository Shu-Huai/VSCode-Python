from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def AddTitle(resultDocument, text, isCenter=True):
    title = resultDocument.add_paragraph()
    if isCenter:
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(text).font.bold = True
    if isCenter:
        resultDocument.add_paragraph()


def AddChoice(resultDocument, paragraphs, numbers):
    AddTitle(resultDocument, paragraphs[0].text, False)
    paragraphs.pop(0)
    for i in range(numbers):
        answers = list(re.findall("[A-Z]", str(paragraphs[0].text)))
        resultDocument.add_paragraph(paragraphs[0].text)
        paragraphs.pop(0)
        temp = 0
        selections = ""
        thisLine = str(paragraphs[0].text)
        while len(thisLine) and thisLine[0] != "【":
            selections += thisLine
            temp += 1
            thisLine = str(paragraphs[temp].text)
        for i in range(temp):
            paragraphs.pop(0)
        selections = list(re.findall("[A-Z][^A-Z]+", selections))
        for i in range(len(selections)):
            temp = selections[i].strip()
            newParagraph = resultDocument.add_paragraph()
            if temp[0] in answers:
                newRun = newParagraph.add_run(temp)
                newRun.font.bold = True
                newRun.font.color.rgb = RGBColor(255, 0, 0)
            else:
                newParagraph.add_run(temp)
    resultDocument.add_paragraph()
    paragraphs.pop(0)


def AddJudge(resultDocument, paragraphs, numbers):
    AddTitle(resultDocument, paragraphs[0].text, False)
    paragraphs.pop(0)
    for i in range(numbers):
        resultDocument.add_paragraph(paragraphs[0].text)
        paragraphs.pop(0)


sourceDocument = Document(r"Docx Auto Bold\Source Document.docx")
resultDocument = Document()
resultDocument.styles["Normal"].font.name = u"宋体"
resultDocument.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
resultDocument.styles["Normal"].font.size = Pt(12)
resultDocument.styles["Normal"].font.name = "Times New Roman"
paragraphs = list(sourceDocument.paragraphs)
AddTitle(resultDocument, paragraphs[0].text)
paragraphs.pop(0)
paragraphs.pop(0)
AddChoice(resultDocument, paragraphs, 248)
AddChoice(resultDocument, paragraphs, 208)
AddJudge(resultDocument, paragraphs, 152)
resultDocument.save(r"Docx Auto Bold\Result Document.docx")