import sys
import os
from PyQt5.QtWidgets import QDialog, QLabel, QPushButton, QApplication, QMainWindow
from PyQt5 import QtWidgets
from PyQt5.QtCore import QRect
from PyQt5.QtGui import QFont
from PyQt5.QtGui import QIcon
import docx
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from urllib import request

from docx.image.image import Image
from docx.shared import Inches


@property
def image_width(self):
    if (self.horz_dpi == 0):
        return Inches(self.px_width / 72)
    return Inches(self.px_width / self.horz_dpi)


@property
def image_height(self):
    if (self.vert_dpi == 0):
        return Inches(self.px_height / 72)
    return Inches(self.px_height / self.vert_dpi)


Image.width = image_width
Image.height = image_height


class SteamDiscountInformationGetter(object):
    def __init__(self) -> None:
        super().__init__()

    def GetMaxPage(self):
        url = []
        url.append("https://store.steampowered.com/search/?specials=1&page=1")
        soup = self.GetSoup(self.GetUrlContents(url))
        node = soup[0].find_all("div", class_="search_pagination_right")
        return int(node[0].contents[5].contents[0])

    def CreateUrls(pages):
        urls = []
        for i in range(pages):
            urlExample = "https://store.steampowered.com/search/?specials=1&page={}".format(i + 1)
            urls.append(urlExample)
        return urls

    def GetUrlContents(urls):
        headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_16_6) AppleWebKit/605.1.15 (KHTML, like Gecko) Chrome/ 90.0.4430.93 Safari/605.1.15'}
        responseList = []
        contentList = []
        for i in range(len(urls)):
            try:
                response = requests.get(urls[i], headers=headers)
            except requests.exceptions.ConnectionError:
                MainWindow.EjectConnectionErrorDialog(MainWindow)
                exit()
            responseList.append(response)
            contentList.append(responseList[i].text)
        return contentList

    def GetSoup(contentList):
        soup = []
        for i in range(len(contentList)):
            soup.append(BeautifulSoup(contentList[i], "html.parser"))
        return soup

    def GetGameNames(self, contentList):
        names = []
        soup = self.GetSoup(contentList)
        for i in range(len(contentList)):
            names.extend(soup[i].find_all("span", class_="title"))
        for i in range(len(names)):
            names[i] = names[i].string
        return names

    def GetGameUrls(self, contentList):
        urls = []
        soup = self.GetSoup(contentList)
        urlPrefix = "https://store.steampowered.com/"
        for i in range(len(contentList)):
            for node in soup[i].find_all("a"):
                temp = node.get("href")
                if (urlPrefix + "app/" in temp or urlPrefix + "bundle/" in temp or urlPrefix + "sub/" in temp) and "view" not in temp:
                    urls.append(node.get("href"))
        return urls

    def GetPrices(self, contentList):
        previousPrices = []
        nowPrices = []
        discounts = []
        soup = self.GetSoup(contentList)
        for i in range(len(contentList)):
            count = 0
            unpurchaseableIndex = []
            for node in soup[i].find_all("div", class_="col search_discount responsive_secondrow"):
                discount = node.text.strip("\n")
                if discount == "":
                    discounts.append("0")
                    unpurchaseableIndex.append(count)
                else:
                    discounts.append(discount)
                count += 1
            for node in soup[i].find_all("div", class_="col search_price discounted responsive_secondrow"):
                previousPrices.append(node.contents[1].contents[0].contents[0])
                nowPrices.append(node.contents[3].strip())
            for j in unpurchaseableIndex:
                previousPrices.insert(i * 25 + j, "Unpurchasable")
                nowPrices.insert(i * 25 + j, "Unpurchasable")
        return previousPrices, nowPrices, discounts

    def DeleteGameCovers(self, path):
        for i in os.listdir(path):
            file = path + "\\" + i
            if os.path.isfile(file):
                os.remove(file)
            else:
                self.DeleteGameCovers(file)

    def GetGameCovers(self, contentList):
        if not os.path.exists(r"Steam Discount Information Getter\Game Cover"):
            os.makedirs(r"Steam Discount Information Getter\Game Cover")
        self.DeleteGameCovers(self, r"Steam Discount Information Getter\Game Cover")
        soup = self.GetSoup(contentList)
        count = 0
        for i in range(len(contentList)):
            for node in soup[i].find_all("div", class_="col search_capsule"):
                url = node.contents[0].attrs["src"].replace("capsule_sm_120.jpg", "header.jpg")
                request.urlretrieve(url, r"Steam Discount Information Getter\Game Cover" + "\\" + str(count) + ".png")
                count += 1

    def Merge(names, urls, previousPrices, nowPrices, discounts):
        games = []
        for i in range(len(names)):
            dictionary = dict(gameName=names[i], gameUrl=urls[i], previousPrice=previousPrices[i], nowPrice=nowPrices[i], discount=discounts[i], gameCoverNumber=i)
            games.append(dictionary)
        return games

    def Sort(games, sortRule):
        for i in range(len(games)):
            if games[i]["previousPrice"] == "Unpurchasable":
                games[i]["previousPrice"] = "1000000"
            if games[i]["nowPrice"] == "Unpurchasable":
                games[i]["nowPrice"] = "1000000"
            games[i]["previousPrice"] = float(games[i]["previousPrice"].strip("¥"))
            games[i]["nowPrice"] = float(games[i]["nowPrice"].strip("¥"))
            games[i]["discount"] = float(games[i]["discount"].strip("%"))
        games = sorted(games, key=lambda x: (x[sortRule], x["gameName"]))
        for i in range(len(games)):
            if games[i]["previousPrice"] == 1000000:
                games[i]["previousPrice"] = "Unpurchasable"
            else:
                games[i]["previousPrice"] = "¥ " + str(games[i]["previousPrice"])
            if games[i]["nowPrice"] == 1000000:
                games[i]["nowPrice"] = "Unpurchasable"
            else:
                games[i]["nowPrice"] = "¥ " + str(games[i]["nowPrice"])
            games[i]["discount"] = str(games[i]["discount"]) + "%"
        return games

    def SaveToDocx(games):
        docxFile = Document()
        docxFile.styles["Normal"].font.name = u"宋体"
        docxFile.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        docxFile.styles["Normal"].font.size = Pt(12)
        docxFile.add_paragraph().add_run("Steam平台优惠信息").font.bold = True
        for i in range(len(games)):
            paragraph = docxFile.add_paragraph()
            paragraph.add_run("游戏：%s.\n" % games[i]["gameName"]).font.bold = True
            paragraph.add_run("链接：")
            part = paragraph.part
            ralationId = part.relate_to(games[i]["gameUrl"], docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            hyperLink = docx.oxml.shared.OxmlElement("w:hyperlink")
            hyperLink.set(
                docx.oxml.shared.qn("r:id"),
                ralationId,
            )
            r = docx.oxml.shared.OxmlElement('w:r')
            rPr = docx.oxml.shared.OxmlElement('w:rPr')
            r.append(rPr)
            r.text = games[i]["gameUrl"]
            hyperLink.append(r)
            run = paragraph.add_run()
            run._r.append(hyperLink)
            run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
            run.font.underline = True
            paragraph.add_run("\n折扣：%s，" % games[i]["discount"])
            paragraph.add_run("价格：%s，前价%s。" % (games[i]["nowPrice"], games[i]["previousPrice"]))
            docxFile.add_picture(r"Steam Discount Information Getter\Game Cover" + "\\" + str(games[i]["gameCoverNumber"]) + ".png")
        docxFile.save(r"Steam Discount Information Getter\Steam Discount Information.docx")


class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.setWindowTitle("Steam Discount Information Getter")
        self.setWindowIcon(QIcon(r"Steam Discount Information Getter\Steam Discount Information Getter.png"))
        self.setObjectName("mainWindow")
        self.resize(1500, 500)
        self.centralWidget = QtWidgets.QWidget(self)
        self.centralWidget.setObjectName("centralWidget")
        self.InitializeTitleLabel()
        self.InitializePageNumberLabel()
        self.InitializePageNumberEdit()
        self.InitializeOkButton()
        self.InitializeCloseButton()
        self.InitializeSortRuleGroup()
        self.InitializeSaveCheckBox()
        self.InitializeResultTextBrowser()
        self.InitializeClearButton()
        self.pageNumberLabel.raise_()
        self.okButton.raise_()
        self.titleLabel.raise_()
        self.pageNumberEdit.raise_()
        self.closeButton.raise_()
        self.setCentralWidget(self.centralWidget)
        self.menubar = QtWidgets.QMenuBar(self)
        self.menubar.setGeometry(QRect(0, 0, 800, 22))
        self.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(self)
        self.setStatusBar(self.statusbar)

    def InitializeTitleLabel(self):
        self.titleLabel = QtWidgets.QLabel("Steam平台折扣信息爬取", self.centralWidget)
        self.titleLabel.setGeometry(QRect(170, 50, 391, 40))
        self.titleLabel.setFont(QFont("宋体", 20))

    def InitializePageNumberLabel(self):
        pages = SteamDiscountInformationGetter.GetMaxPage(SteamDiscountInformationGetter)
        self.pageNumberLabel = QtWidgets.QLabel("请输入要爬取的信息页数：        /" + str(pages), self.centralWidget)
        self.pageNumberLabel.setGeometry(QRect(170, 110, 301, 21))
        self.pageNumberLabel.setFont(QFont("宋体", 12))

    def InitializePageNumberEdit(self):
        self.pageNumberEdit = QtWidgets.QLineEdit(self.centralWidget)
        self.pageNumberEdit.setGeometry(QRect(370, 110, 41, 21))
        self.pageNumberEdit.setFont(QFont("宋体", 12))
        self.pageNumberEdit.returnPressed.connect(self.OkButtonClicked)

    def InitializeOkButton(self):
        self.okButton = QtWidgets.QPushButton("确认", self.centralWidget)
        self.okButton.setGeometry(QRect(390, 310, 75, 24))
        self.okButton.setFont(QFont("宋体", 12))
        self.okButton.clicked.connect(self.OkButtonClicked)

    def InitializeCloseButton(self):
        self.closeButton = QtWidgets.QPushButton("关闭", self.centralWidget)
        self.closeButton.setGeometry(QRect(480, 310, 75, 24))
        self.closeButton.setFont(QFont("宋体", 12))
        self.closeButton.clicked.connect(self.close)

    def InitializeSortRuleGroup(self):
        self.sortRuleGroup = QtWidgets.QGroupBox("排序规则：", self.centralWidget)
        self.sortRuleGroup.setGeometry(QRect(170, 160, 391, 80))
        self.sortRuleGroup.setFont(QFont("宋体", 12))
        self.InitializeGameNameRadio()
        self.InitializePreviousPriceRadio()
        self.InitializeNowPriceRadio()
        self.InitializeDiscountRadio()
        self.InitializeGameCoverNumberRadio()

    def InitializeGameNameRadio(self):
        self.gameNameRadio = QtWidgets.QRadioButton("游戏名称", self.sortRuleGroup)
        self.gameNameRadio.setGeometry(QRect(10, 20, 95, 20))
        self.gameNameRadio.setFont(QFont("宋体", 12))
        self.gameNameRadio.clicked.connect(self.GameNameRadioChecked)

    def InitializePreviousPriceRadio(self):
        self.previousPriceRadio = QtWidgets.QRadioButton("先前价格", self.sortRuleGroup)
        self.previousPriceRadio.setGeometry(QRect(140, 20, 111, 20))
        self.previousPriceRadio.setFont(QFont("宋体", 12))
        self.previousPriceRadio.clicked.connect(self.PreviousPriceRadioChecked)

    def InitializeNowPriceRadio(self):
        self.nowPriceRadio = QtWidgets.QRadioButton("当前价格", self.sortRuleGroup)
        self.nowPriceRadio.setGeometry(QRect(270, 20, 95, 20))
        self.nowPriceRadio.setFont(QFont("宋体", 12))
        self.nowPriceRadio.clicked.connect(self.NowPriceRadioChecked)

    def InitializeDiscountRadio(self):
        self.discountRadio = QtWidgets.QRadioButton("折扣力度", self.sortRuleGroup)
        self.discountRadio.setGeometry(QRect(10, 50, 95, 20))
        self.discountRadio.setFont(QFont("宋体", 12))
        self.discountRadio.setChecked(True)
        self.discountRadio.clicked.connect(self.DiscountRadioChecked)

    def InitializeGameCoverNumberRadio(self):
        self.gameCoverNumberRadio = QtWidgets.QRadioButton("默认顺序", self.sortRuleGroup)
        self.gameCoverNumberRadio.setGeometry(QRect(140, 50, 171, 20))
        self.gameCoverNumberRadio.setFont(QFont("宋体", 12))
        self.gameCoverNumberRadio.clicked.connect(self.GameCoverNumberRadioChecked)

    def InitializeSaveCheckBox(self):
        self.saveCheckBox = QtWidgets.QCheckBox("保存数据到文档", self.centralWidget)
        self.saveCheckBox.setGeometry(QRect(180, 260, 221, 20))
        self.saveCheckBox.setFont(QFont("宋体", 12))
        self.saveCheckBox.setChecked(True)

    def InitializeResultTextBrowser(self):
        self.resultTextBrowser = QtWidgets.QTextBrowser(self.centralWidget)
        self.resultTextBrowser.setGeometry(QRect(630, 60, 791, 271))
        self.resultTextBrowser.setFont(QFont("宋体", 12))
        self.resultTextBrowser.setOpenExternalLinks(True)

    def InitializeClearButton(self):
        self.clearButton = QtWidgets.QPushButton("清除", self.centralWidget)
        self.clearButton.setGeometry(QRect(630, 350, 75, 24))
        self.clearButton.setObjectName("clearButton")
        self.clearButton.setFont(QFont("宋体", 12))
        self.clearButton.clicked.connect(self.resultTextBrowser.clear)

    def GetSortRule(self):
        if self.gameNameRadio.isChecked():
            return "gameName"
        if self.previousPriceRadio.isChecked():
            return "previousPrice"
        if self.nowPriceRadio.isChecked():
            return "nowPrice"
        if self.discountRadio.isChecked():
            return "discount"
        if self.gameCoverNumberRadio.isChecked():
            return "gameCoverNumber"

    def EjectConnectionErrorDialog(self):
        errorDialog = QDialog()
        errorDialog.resize(260, 100)
        errorDialog.setWindowTitle("错误")
        errorDialog.setWindowIcon(QIcon(r"Steam Discount Information Getter\Steam Discount Information Getter Error.png"))
        errorLabel = QLabel("网络连接错误。\n请检查网络连接。", errorDialog)
        errorLabel.move(30, 10)
        errorLabel.setFont(QFont("宋体", 12))
        okButton = QPushButton("确认", errorDialog)
        okButton.move(110, 60)
        okButton.setFont(QFont("宋体", 12))
        okButton.clicked.connect(errorDialog.close)
        errorDialog.exec_()

    def EjectPageNumberErrorDialog(self):
        errorDialog = QDialog()
        errorDialog.resize(180, 100)
        errorDialog.setWindowTitle("错误")
        errorDialog.setWindowIcon(QIcon(r"Steam Discount Information Getter\Steam Discount Information Getter Error.png"))
        errorLabel = QLabel("错误的页数。\n请重新输入。", errorDialog)
        errorLabel.move(30, 10)
        errorLabel.setFont(QFont("宋体", 12))
        okButton = QPushButton("确认", errorDialog)
        okButton.move(50, 60)
        okButton.setFont(QFont("宋体", 12))
        okButton.clicked.connect(errorDialog.close)
        errorDialog.exec_()

    def OkButtonClicked(self):
        if self.pageNumberEdit.text() == "":
            self.pageNumberEdit.setText("5")
        try:
            pages = int(self.pageNumberEdit.text())
        except ValueError:
            self.EjectPageNumberErrorDialog()
            self.pageNumberEdit.clear()
            self.pageNumberEdit.setFocus()
            return
        if pages < 0 or pages > int(self.pageNumberLabel.text().strip("请输入要爬取的信息页数：        /")):
            self.EjectPageNumberErrorDialog()
            self.pageNumberEdit.clear()
            self.pageNumberEdit.setFocus()
            return
        urls = SteamDiscountInformationGetter.CreateUrls(pages)
        contentList = SteamDiscountInformationGetter.GetUrlContents(urls)
        gameNames = SteamDiscountInformationGetter.GetGameNames(SteamDiscountInformationGetter, contentList)
        gameUrls = SteamDiscountInformationGetter.GetGameUrls(SteamDiscountInformationGetter, contentList)
        previousPrices, nowPrices, discounts = SteamDiscountInformationGetter.GetPrices(SteamDiscountInformationGetter, contentList)
        SteamDiscountInformationGetter.GetGameCovers(SteamDiscountInformationGetter, contentList)
        games = SteamDiscountInformationGetter.Merge(gameNames, gameUrls, previousPrices, nowPrices, discounts)
        games = SteamDiscountInformationGetter.Sort(games, sortRule=self.GetSortRule())
        self.resultTextBrowser.clear()
        for i in range(len(games)):
            self.resultTextBrowser.append("游戏%s。" % (games[i]["gameName"]))
            self.resultTextBrowser.append('链接：<a href=%s>%s</a>。' % (games[i]["gameUrl"], games[i]["gameUrl"]))
            self.resultTextBrowser.append("折扣：%s，价格：%s，前价%s。\n" % (games[i]["discount"], games[i]["nowPrice"], games[i]["previousPrice"]))
            self.resultTextBrowser.insertHtml('<img src="Steam Discount Information Getter\Game Cover\%d.png"/>' % games[i]["gameCoverNumber"])
            self.resultTextBrowser.append("")
        if self.saveCheckBox.isChecked():
            SteamDiscountInformationGetter.SaveToDocx(games)
        completeDialog = QDialog()
        completeDialog.resize(180, 100)
        completeDialog.setWindowTitle("成功")
        completeDialog.setWindowIcon(QIcon(r"Steam Discount Information Getter\Steam Discount Information Getter Complete.png"))
        completeLabel = QLabel("操作成功。", completeDialog)
        completeLabel.move(40, 20)
        completeLabel.setFont(QFont("宋体", 12))
        okButton = QPushButton("确认", completeDialog)
        okButton.move(50, 60)
        okButton.setFont(QFont("宋体", 12))
        okButton.clicked.connect(completeDialog.close)
        completeDialog.exec_()

    def GameNameRadioChecked(self):
        self.previousPriceRadio.setChecked(False)
        self.nowPriceRadio.setChecked(False)
        self.discountRadio.setChecked(False)
        self.gameCoverNumberRadio.setChecked(False)

    def PreviousPriceRadioChecked(self):
        self.gameNameRadio.setChecked(False)
        self.nowPriceRadio.setChecked(False)
        self.discountRadio.setChecked(False)
        self.gameCoverNumberRadio.setChecked(False)

    def NowPriceRadioChecked(self):
        self.gameNameRadio.setChecked(False)
        self.previousPriceRadio.setChecked(False)
        self.discountRadio.setChecked(False)
        self.gameCoverNumberRadio.setChecked(False)

    def DiscountRadioChecked(self):
        self.gameNameRadio.setChecked(False)
        self.previousPriceRadio.setChecked(False)
        self.nowPriceRadio.setChecked(False)
        self.gameCoverNumberRadio.setChecked(False)

    def GameCoverNumberRadioChecked(self):
        self.gameNameRadio.setChecked(False)
        self.previousPriceRadio.setChecked(False)
        self.nowPriceRadio.setChecked(False)
        self.discountRadio.setChecked(False)


application = QApplication(sys.argv)
mainWindow = MainWindow()
mainWindow.show()
sys.exit(application.exec_())